from pathlib import Path

from docx import Document
from docx.enum.table import WD_CELL_VERTICAL_ALIGNMENT, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


OUT_DOCX = Path("loan_contract_preview.docx")
OUT_HTML = Path("loan_contract_preview.html")


def set_run_font(run, size=None, bold=None, italic=None, color=None, name="Calibri"):
    run.font.name = name
    run._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    if size is not None:
        run.font.size = Pt(size)
    if bold is not None:
        run.bold = bold
    if italic is not None:
        run.italic = italic
    if color is not None:
        run.font.color.rgb = color


def set_cell_text(cell, text, *, bold=False, size=10.5, color=None, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    p.paragraph_format.line_spacing = 1.0
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(0)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, color=color)
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER


def shade_cell(cell, fill):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_margins(cell, top=80, start=120, bottom=80, end=120):
    tc_pr = cell._tc.get_or_add_tcPr()
    tc_mar = tc_pr.first_child_found_in("w:tcMar")
    if tc_mar is None:
        tc_mar = OxmlElement("w:tcMar")
        tc_pr.append(tc_mar)
    for name, value in [("top", top), ("start", start), ("bottom", bottom), ("end", end)]:
        node = tc_mar.find(qn(f"w:{name}"))
        if node is None:
            node = OxmlElement(f"w:{name}")
            tc_mar.append(node)
        node.set(qn("w:w"), str(value))
        node.set(qn("w:type"), "dxa")


def set_table_widths(table, widths_dxa):
    table.autofit = False
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    tbl = table._tbl
    tbl_pr = tbl.tblPr
    tbl_w = tbl_pr.first_child_found_in("w:tblW")
    if tbl_w is None:
        tbl_w = OxmlElement("w:tblW")
        tbl_pr.append(tbl_w)
    tbl_w.set(qn("w:type"), "dxa")
    tbl_w.set(qn("w:w"), str(sum(widths_dxa)))

    tbl_grid = tbl.tblGrid
    if tbl_grid is None:
        tbl_grid = OxmlElement("w:tblGrid")
        tbl.append(tbl_grid)
    while len(tbl_grid) > 0:
        tbl_grid.remove(tbl_grid[0])
    for width in widths_dxa:
        grid_col = OxmlElement("w:gridCol")
        grid_col.set(qn("w:w"), str(width))
        tbl_grid.append(grid_col)

    for row in table.rows:
        for idx, cell in enumerate(row.cells):
            tc_pr = cell._tc.get_or_add_tcPr()
            tc_w = tc_pr.first_child_found_in("w:tcW")
            if tc_w is None:
                tc_w = OxmlElement("w:tcW")
                tc_pr.append(tc_w)
            tc_w.set(qn("w:type"), "dxa")
            tc_w.set(qn("w:w"), str(widths_dxa[idx]))
            set_cell_margins(cell)


def add_para(doc, text="", *, style=None, align=None, size=None, bold=None, italic=None, color=None, before=None, after=None, line=1.15):
    p = doc.add_paragraph(style=style)
    if align is not None:
        p.alignment = align
    pf = p.paragraph_format
    pf.line_spacing = line
    if before is not None:
        pf.space_before = Pt(before)
    if after is not None:
        pf.space_after = Pt(after)
    run = p.add_run(text)
    set_run_font(run, size=size, bold=bold, italic=italic, color=color)
    return p


def add_note_box(doc, text):
    table = doc.add_table(rows=1, cols=1)
    set_table_widths(table, [9360])
    cell = table.cell(0, 0)
    shade_cell(cell, "F4F7FB")
    set_cell_text(cell, text, size=10.5, color=RGBColor(77, 77, 77))
    return table


def add_summary_table(doc, rows):
    table = doc.add_table(rows=len(rows), cols=4)
    set_table_widths(table, [1500, 3180, 1500, 3180])
    table.style = "Table Grid"
    for r, (l1, v1, l2, v2) in enumerate(rows):
        cells = table.rows[r].cells
        for c in (0, 2):
            shade_cell(cells[c], "E8EEF5")
        set_cell_text(cells[0], l1, bold=True, size=10, color=RGBColor(31, 77, 120), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[1], v1, size=10.5)
        set_cell_text(cells[2], l2, bold=True, size=10, color=RGBColor(31, 77, 120), align=WD_ALIGN_PARAGRAPH.CENTER)
        set_cell_text(cells[3], v2, size=10.5)
    return table


def add_clause(doc, title, body):
    add_para(doc, title, style="Heading 1", before=10, after=4, line=1.0)
    add_para(doc, body, before=0, after=6, line=1.25)


def set_footer(section, text):
    p = section.footer.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.text = ""
    r = p.add_run(text)
    set_run_font(r, size=8.5, color=RGBColor(118, 118, 118), italic=True)


def build_docx():
    doc = Document()
    section = doc.sections[0]
    section.page_width = Inches(8.5)
    section.page_height = Inches(11)
    for side in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, side, Inches(1))
    section.header_distance = Inches(0.35)
    section.footer_distance = Inches(0.45)

    normal = doc.styles["Normal"]
    normal.font.name = "Calibri"
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
    normal.font.size = Pt(11)

    for style_name, size, color, bold, before, after in [
        ("Title", 20, RGBColor(0, 0, 0), True, 0, 6),
        ("Heading 1", 15, RGBColor(46, 116, 181), True, 12, 6),
        ("Heading 2", 12, RGBColor(31, 77, 120), True, 8, 4),
    ]:
        style = doc.styles[style_name]
        style.font.name = "Calibri"
        style._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft JhengHei")
        style.font.size = Pt(size)
        style.font.bold = bold
        style.font.color.rgb = color
        pf = style.paragraph_format
        pf.space_before = Pt(before)
        pf.space_after = Pt(after)

    set_footer(section, "本文件為系統產生之示意預覽，實際契約內容以核准版本與法務文件為準。")

    add_para(doc, "示意預覽 / 系統產生", align=WD_ALIGN_PARAGRAPH.RIGHT, size=8.5, italic=True, color=RGBColor(120, 120, 120), after=2, line=1.0)
    add_para(doc, "貸款契約書", style="Title", align=WD_ALIGN_PARAGRAPH.CENTER, size=20, bold=True, after=0, line=1.0)
    add_para(doc, "個人信貸契約示意稿", align=WD_ALIGN_PARAGRAPH.CENTER, size=11, italic=True, color=RGBColor(90, 90, 90), after=10, line=1.0)
    add_note_box(doc, "本文件為根據填單與審核資料生成之契約示意預覽，主要用於版面與欄位確認；正式文件將依實際核准結果、法務條款與簽核流程另行定稿。")

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    add_para(doc, "一、契約條件摘要", style="Heading 1", before=14, after=6, line=1.0)
    add_summary_table(doc, [
        ("契約編號", "LC-20260520-001", "申請編號", "LA202605200001"),
        ("借款人", "王大明", "身分證字號", "A123456789"),
        ("貸款類型", "個人信貸", "貸款用途", "個人資金週轉"),
        ("核准金額", "NT$ 13,000", "期數", "12 期"),
        ("年利率", "4.00%", "每月應繳", "NT$ 1,107"),
        ("貸款帳號", "LAC202605190004301909", "撥款帳號", "070000000027"),
        ("首期繳款日", "2027-05-19", "簽約日期", "中華民國115年5月20日"),
        ("聯絡電話", "0912-345-678", "電子郵件", "chiu890909@gmail.com"),
    ])

    add_para(doc, "二、契約主文", style="Heading 1", before=14, after=4, line=1.0)
    add_para(doc, "本契約由借款人王大明（以下稱「借款人」）與 Java Easy Bank（以下稱「本行」）依雙方約定條件簽訂。借款人同意依本契約及系統所載條件，向本行申請並使用本筆貸款。", before=0, after=6, line=1.25)

    add_clause(doc, "第1條　貸款金額與期間", "本行同意於核准後提供借款人新臺幣 13,000 元整之貸款額度，還款期數為 12 期。借款人應依本契約及系統通知之還款計畫按期清償。")
    add_clause(doc, "第2條　撥款與還款方式", "本行得於核准條件完成後，依約將款項撥入借款人指定帳戶。借款人應自首期繳款日起，按月依系統顯示之應繳金額扣款或繳納；若因四捨五入或提前清償產生差額，以實際結算金額為準。")
    add_clause(doc, "第3條　利率與費用", "本貸款年利率為 4.00%。除本契約另有約定外，借款人應依實際動用本金及期間計息；如有逾期、補費或其他費用，依本行公告、系統提示及相關規定辦理。")
    add_clause(doc, "第4條　提前清償", "借款人得於貸款存續期間內申請提前清償。本行得依清償日之實際餘額、未到期利息及相關費用計算最終應付金額，借款人同意以系統結算結果為準。")
    add_clause(doc, "第5條　遲延與違約", "借款人如未依約於應繳日前完成繳款，本行得依本契約及相關規定進行催收、收取遲延費用，並採取必要之風險控管措施；如情節重大，本行得依約提前終止本契約。")
    add_clause(doc, "第6條　資料真實性與通知", "借款人聲明本次申請與填載資料均屬真實、完整且無重大遺漏。借款人同意本行得以電子郵件、簡訊或系統通知方式送達本契約相關訊息，並視為合法有效之通知。")
    add_clause(doc, "第7條　其他約定", "本契約未盡事宜，悉依中華民國相關法令、本行作業規則及雙方後續書面或電子約定辦理。若本示意文件與正式核准版本不一致，應以正式版本及系統紀錄為準。")

    add_para(doc, "三、簽署欄位", style="Heading 1", before=14, after=6, line=1.0)
    sig = doc.add_table(rows=2, cols=2)
    set_table_widths(sig, [4680, 4680])
    sig.style = "Table Grid"
    shade_cell(sig.cell(0, 0), "F6F7F8")
    shade_cell(sig.cell(0, 1), "F6F7F8")
    set_cell_text(sig.cell(0, 0), "借款人簽章", bold=True, size=10.5, color=RGBColor(31, 77, 120), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(sig.cell(0, 1), "本行代表簽章", bold=True, size=10.5, color=RGBColor(31, 77, 120), align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(sig.cell(1, 0), "____________________________\n王大明\n日期：__________________", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_text(sig.cell(1, 1), "____________________________\nJava Easy Bank\n日期：__________________", size=10.5, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_para(doc, "備註：本契約為示意預覽版，若您要，我可以再把它改成更像正式銀行版的「精簡條款版」或「法務版」格式。", size=9.5, color=RGBColor(100, 100, 100), italic=True, before=10, after=0, line=1.1)
    doc.save(OUT_DOCX)


def build_html():
    html = """<!DOCTYPE html>
<html lang="zh-Hant">
<head>
  <meta charset="UTF-8" />
  <meta name="viewport" content="width=device-width, initial-scale=1.0" />
  <title>貸款契約書（示意預覽）</title>
  <style>
    :root {
      --ink: #222;
      --muted: #666;
      --line: #d9dee6;
      --panel: #f4f7fb;
      --accent: #2e74b5;
      --soft: #e8eef5;
      --page: #fff;
      --bg: #eef2f7;
    }
    * { box-sizing: border-box; }
    body {
      margin: 0;
      background: var(--bg);
      color: var(--ink);
      font-family: Calibri, "Microsoft JhengHei", Arial, sans-serif;
      line-height: 1.4;
    }
    .page {
      width: min(900px, calc(100vw - 40px));
      margin: 24px auto;
      background: var(--page);
      border: 1px solid var(--line);
      box-shadow: 0 10px 32px rgba(25, 35, 52, 0.08);
      padding: 44px 52px 36px;
    }
    .preview-tag {
      text-align: right;
      color: #7a7a7a;
      font-size: 12px;
      font-style: italic;
      margin-bottom: 6px;
    }
    h1 {
      margin: 0;
      text-align: center;
      font-size: 30px;
      letter-spacing: 0;
      color: #000;
    }
    .subtitle {
      text-align: center;
      margin: 4px 0 14px;
      color: var(--muted);
      font-style: italic;
      font-size: 14px;
    }
    .note {
      background: var(--panel);
      border-left: 4px solid var(--accent);
      padding: 14px 16px;
      margin: 14px 0 18px;
      color: #4d4d4d;
      font-size: 14px;
    }
    h2 {
      margin: 18px 0 10px;
      font-size: 19px;
      color: var(--accent);
      border-bottom: 1px solid #edf1f6;
      padding-bottom: 4px;
    }
    .summary {
      width: 100%;
      border-collapse: collapse;
      table-layout: fixed;
      margin-bottom: 12px;
    }
    .summary th,
    .summary td {
      border: 1px solid var(--line);
      padding: 11px 12px;
      vertical-align: middle;
      word-wrap: break-word;
    }
    .summary th {
      width: 16%;
      background: var(--soft);
      color: #1f4d78;
      font-weight: 700;
      text-align: center;
      font-size: 13px;
    }
    .summary td { width: 34%; font-size: 13.5px; }
    .article {
      margin-top: 10px;
      font-size: 14px;
    }
    .article-title {
      font-weight: 700;
      color: #1f4d78;
      margin: 14px 0 4px;
      font-size: 15px;
    }
    .article p { margin: 0 0 8px; }
    .signatures {
      width: 100%;
      border-collapse: collapse;
      table-layout: fixed;
      margin-top: 12px;
    }
    .signatures th,
    .signatures td {
      border: 1px solid var(--line);
      padding: 12px;
      text-align: center;
      vertical-align: middle;
    }
    .signatures th {
      background: #f6f7f8;
      color: #1f4d78;
      width: 50%;
    }
    .signature-line {
      margin-top: 10px;
      font-family: "Courier New", monospace;
      letter-spacing: 0.3px;
    }
    .footer-note {
      margin-top: 12px;
      color: #666;
      font-size: 12.5px;
      font-style: italic;
    }
    @media print {
      body { background: #fff; }
      .page {
        width: auto;
        margin: 0;
        border: 0;
        box-shadow: none;
        padding: 24px 28px;
      }
    }
  </style>
</head>
<body>
  <div class="page">
    <div class="preview-tag">示意預覽 / 系統產生</div>
    <h1>貸款契約書</h1>
    <div class="subtitle">個人信貸契約示意稿</div>
    <div class="note">本文件為根據填單與審核資料生成之契約示意預覽，主要用於版面與欄位確認；正式文件將依實際核准結果、法務條款與簽核流程另行定稿。</div>
    <h2>一、契約條件摘要</h2>
    <table class="summary">
      <tr><th>契約編號</th><td>LC-20260520-001</td><th>申請編號</th><td>LA202605200001</td></tr>
      <tr><th>借款人</th><td>王大明</td><th>身分證字號</th><td>A123456789</td></tr>
      <tr><th>貸款類型</th><td>個人信貸</td><th>貸款用途</th><td>個人資金週轉</td></tr>
      <tr><th>核准金額</th><td>NT$ 13,000</td><th>期數</th><td>12 期</td></tr>
      <tr><th>年利率</th><td>4.00%</td><th>每月應繳</th><td>NT$ 1,107</td></tr>
      <tr><th>貸款帳號</th><td>LAC202605190004301909</td><th>撥款帳號</th><td>070000000027</td></tr>
      <tr><th>首期繳款日</th><td>2027-05-19</td><th>簽約日期</th><td>中華民國115年5月20日</td></tr>
      <tr><th>聯絡電話</th><td>0912-345-678</td><th>電子郵件</th><td>chiu890909@gmail.com</td></tr>
    </table>
    <h2>二、契約主文</h2>
    <div class="article">
      <p>本契約由借款人王大明（以下稱「借款人」）與 Java Easy Bank（以下稱「本行」）依雙方約定條件簽訂。借款人同意依本契約及系統所載條件，向本行申請並使用本筆貸款。</p>
      <div class="article-title">第1條　貸款金額與期間</div>
      <p>本行同意於核准後提供借款人新臺幣 13,000 元整之貸款額度，還款期數為 12 期。借款人應依本契約及系統通知之還款計畫按期清償。</p>
      <div class="article-title">第2條　撥款與還款方式</div>
      <p>本行得於核准條件完成後，依約將款項撥入借款人指定帳戶。借款人應自首期繳款日起，按月依系統顯示之應繳金額扣款或繳納；若因四捨五入或提前清償產生差額，以實際結算金額為準。</p>
      <div class="article-title">第3條　利率與費用</div>
      <p>本貸款年利率為 4.00%。除本契約另有約定外，借款人應依實際動用本金及期間計息；如有逾期、補費或其他費用，依本行公告、系統提示及相關規定辦理。</p>
      <div class="article-title">第4條　提前清償</div>
      <p>借款人得於貸款存續期間內申請提前清償。本行得依清償日之實際餘額、未到期利息及相關費用計算最終應付金額，借款人同意以系統結算結果為準。</p>
      <div class="article-title">第5條　遲延與違約</div>
      <p>借款人如未依約於應繳日前完成繳款，本行得依本契約及相關規定進行催收、收取遲延費用，並採取必要之風險控管措施；如情節重大，本行得依約提前終止本契約。</p>
      <div class="article-title">第6條　資料真實性與通知</div>
      <p>借款人聲明本次申請與填載資料均屬真實、完整且無重大遺漏。借款人同意本行得以電子郵件、簡訊或系統通知方式送達本契約相關訊息，並視為合法有效之通知。</p>
      <div class="article-title">第7條　其他約定</div>
      <p>本契約未盡事宜，悉依中華民國相關法令、本行作業規則及雙方後續書面或電子約定辦理。若本示意文件與正式核准版本不一致，應以正式版本及系統紀錄為準。</p>
    </div>
    <h2>三、簽署欄位</h2>
    <table class="signatures">
      <tr><th>借款人簽章</th><th>本行代表簽章</th></tr>
      <tr>
        <td><div class="signature-line">____________________________</div><div>王大明</div><div>日期：__________________</div></td>
        <td><div class="signature-line">____________________________</div><div>Java Easy Bank</div><div>日期：__________________</div></td>
      </tr>
    </table>
    <div class="footer-note">備註：本契約為示意預覽版，若您要，我可以再把它改成更像正式銀行版的「精簡條款版」或「法務版」格式。</div>
  </div>
</body>
</html>"""
    OUT_HTML.write_text(html, encoding="utf-8")


if __name__ == "__main__":
    build_docx()
    build_html()
    print(OUT_DOCX)
    print(OUT_HTML)
